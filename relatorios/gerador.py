"""
Gerador de rascunho de laudo neuropsicológico em formato Word (.docx).
Estrutura baseada em Malloy-Diniz et al. (2016) e Diniz et al. (Avaliação Neuropsicológica, 2ª ed.).
"""

from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json


CLASSIFICACAO_COR = {
    "Muito Rebaixado":  RGBColor(0xC0, 0x00, 0x00),
    "Rebaixado":        RGBColor(0xFF, 0x00, 0x00),
    "Abaixo da Média":  RGBColor(0xFF, 0x82, 0x00),
    "Médio":            RGBColor(0x00, 0x70, 0x00),
    "Acima da Média":   RGBColor(0x00, 0x70, 0x00),
    "Superior":         RGBColor(0x00, 0x43, 0x8A),
    "Muito Superior":   RGBColor(0x00, 0x43, 0x8A),
    "Mínimo":           RGBColor(0x00, 0x70, 0x00),
    "Leve":             RGBColor(0xFF, 0x82, 0x00),
    "Moderado":         RGBColor(0xFF, 0x00, 0x00),
    "Grave":            RGBColor(0xC0, 0x00, 0x00),
    "Dentro da normalidade": RGBColor(0x00, 0x70, 0x00),
}


def _add_heading(doc, texto, nivel=1):
    p = doc.add_heading(texto, level=nivel)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_par(doc, texto, negrito=False, italico=False, cor=None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    if cor:
        run.font.color.rgb = cor
    return p


def _add_campo(doc, rotulo, valor, cor_valor=None):
    p = doc.add_paragraph()
    run_r = p.add_run(f"{rotulo}: ")
    run_r.bold = True
    run_v = p.add_run(str(valor))
    if cor_valor:
        run_v.font.color.rgb = cor_valor
    return p


def gerar_laudo(paciente, avaliacao, resultados, insights=None):
    """
    paciente:   dict do banco (nome, data_nascimento, escolaridade, sexo, queixa_principal)
    avaliacao:  dict do banco (data_avaliacao, observacoes_comportamentais)
    resultados: list de dicts {teste_codigo, escores_calculados, interpretacao}
    insights:   str opcional com análise integrada da IA
    Retorna BytesIO com o arquivo .docx.
    """
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Cabeçalho ──────────────────────────────────────────────────────────
    titulo = doc.add_heading("LAUDO NEUROPSICOLÓGICO", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 1. Identificação ───────────────────────────────────────────────────
    _add_heading(doc, "1. IDENTIFICAÇÃO", 1)

    _add_campo(doc, "Nome", paciente.get("nome", ""))
    _add_campo(doc, "Data de nascimento", paciente.get("data_nascimento", ""))

    idade = _calcular_idade(paciente.get("data_nascimento", ""))
    if idade:
        _add_campo(doc, "Idade", f"{idade} anos")

    _add_campo(doc, "Escolaridade", paciente.get("escolaridade", ""))
    _add_campo(doc, "Sexo", paciente.get("sexo", ""))
    _add_campo(doc, "Data(s) da avaliação", avaliacao.get("data_avaliacao", str(date.today())))

    # ── 2. Motivo da Avaliação ─────────────────────────────────────────────
    _add_heading(doc, "2. MOTIVO DA AVALIAÇÃO E QUEIXA PRINCIPAL", 1)
    doc.add_paragraph(
        paciente.get("queixa_principal", "[Descrever queixa principal e motivo do encaminhamento]")
    )

    # ── 3. Instrumentos Utilizados ─────────────────────────────────────────
    _add_heading(doc, "3. INSTRUMENTOS UTILIZADOS", 1)
    testes_nomes = {
        "RAVLT":  "Teste de Aprendizagem Auditivo-Verbal de Rey (RAVLT)",
        "FDT":    "Five Digit Test (FDT)",
        "TMT":    "Trail Making Test (TMT) — Partes A e B",
        "SRS2":   "Escala de Responsividade Social — 2ª edição (SRS-2)",
        "ETDAH":  "Escala de Avaliação de TDAH — DSM-5 (ETDAH)",
        "BDI2":   "Inventário de Depressão de Beck — 2ª edição (BDI-II)",
        "BAI":    "Inventário de Ansiedade de Beck (BAI)",
        "WAIS":   "Escala de Inteligência Wechsler para Adultos (WAIS)",
        "WISC":   "Escala de Inteligência Wechsler para Crianças (WISC)",
        "BPA":    "Bateria Psicológica para Avaliação da Atenção (BPA)",
        "BAY":    "Bateria de Avaliação de Yamashita (BAY)",
    }
    codigos_usados = [r["teste_codigo"] for r in resultados]
    for codigo in codigos_usados:
        doc.add_paragraph(
            f"• {testes_nomes.get(codigo, codigo)}",
            style="List Bullet"
        )
    if not codigos_usados:
        doc.add_paragraph("[Listar instrumentos aplicados]")

    # ── 4. Observações Comportamentais ─────────────────────────────────────
    _add_heading(doc, "4. OBSERVAÇÕES COMPORTAMENTAIS", 1)
    obs = avaliacao.get("observacoes_comportamentais", "")
    doc.add_paragraph(
        obs if obs else
        "[Descrever comportamento durante a avaliação: nível de alerta, cooperação, "
        "atenção, ansiedade, estratégias espontâneas, padrão de fala, contato visual, etc.]"
    )

    # ── 5. Resultados por Domínio Cognitivo ────────────────────────────────
    _add_heading(doc, "5. RESULTADOS POR DOMÍNIO COGNITIVO", 1)

    MAPA_DOMINIO = {
        "RAVLT":  ("Memória",              "Memória Episódica Verbal"),
        "FDT":    ("Funções Executivas",    "Velocidade de Processamento / Controle Inibitório / Flexibilidade (FDT)"),
        "TMT":    ("Funções Executivas",    "Velocidade de Processamento / Alternância Atencional (TMT)"),
        "SRS2":   ("Cognição Social",       "Responsividade Social (SRS-2)"),
        "ETDAH":  ("Atenção e FE",          "Sintomas de TDAH — DSM-5 (ETDAH)"),
        "BDI2":   ("Aspectos Emocionais",   "Sintomatologia Depressiva (BDI-II)"),
        "BAI":    ("Aspectos Emocionais",   "Sintomatologia Ansiosa (BAI)"),
    }

    dominios_escritos = set()
    for resultado in resultados:
        codigo  = resultado.get("teste_codigo", "")
        escores = resultado.get("escores_calculados", {})
        interp  = resultado.get("interpretacao", "")
        dominio, titulo_subsecao = MAPA_DOMINIO.get(codigo, ("Outros", codigo))

        if dominio not in dominios_escritos:
            _add_heading(doc, f"5.{len(dominios_escritos)+1}. {dominio}", 2)
            dominios_escritos.add(dominio)

        _add_heading(doc, titulo_subsecao, 3)

        # Escore resumido em destaque
        classif = _extrair_classificacao(codigo, escores)
        if classif:
            cor = CLASSIFICACAO_COR.get(classif)
            _add_campo(doc, "Classificação Geral", classif, cor_valor=cor)

        # Interpretação completa
        if interp:
            for linha in interp.split("\n"):
                if linha.strip():
                    doc.add_paragraph(linha)

        doc.add_paragraph()

    # ── 6. Integração dos Achados ──────────────────────────────────────────
    _add_heading(doc, "6. INTEGRAÇÃO DOS ACHADOS", 1)
    if insights:
        for linha in insights.split("\n"):
            if linha.strip():
                doc.add_paragraph(linha)
    else:
        doc.add_paragraph(
            "[Descrever padrões, dissociações e consistência entre os instrumentos. "
            "Relacionar os achados neuropsicológicos à queixa principal e ao contexto clínico do paciente.]"
        )

    # ── 7. Hipóteses Diagnósticas ──────────────────────────────────────────
    _add_heading(doc, "7. HIPÓTESES DIAGNÓSTICAS", 1)
    doc.add_paragraph(
        "[Listar hipóteses em ordem de probabilidade clínica, com justificativa baseada nos dados. "
        "Usar linguagem de hipótese: 'Os dados são compatíveis com...', 'Os achados sugerem...', "
        "'Não se pode descartar...'. Incluir CID-11 e/ou DSM-5-TR quando aplicável.]"
    )

    # ── 8. Conclusão e Recomendações ───────────────────────────────────────
    _add_heading(doc, "8. CONCLUSÃO E RECOMENDAÇÕES", 1)
    doc.add_paragraph(
        "[Síntese objetiva dos achados. Encaminhamentos sugeridos (psiquiatria, fonoaudiologia, "
        "psicoterapia, intervenções escolares, etc.). Orientações para família/paciente.]"
    )

    doc.add_paragraph()
    _add_par(doc, "_" * 50)
    doc.add_paragraph("Neuropsicóloga")
    doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}")

    # ── Rodapé de confiança ────────────────────────────────────────────────
    doc.add_paragraph()
    p_nota = doc.add_paragraph(
        "⚠️ Rascunho gerado automaticamente pela Plataforma NeuroPsico Avaliação. "
        "Os escores e interpretações devem ser revisados e validados pelo(a) neuropsicólogo(a) "
        "responsável antes da emissão do laudo final. "
        "Normas utilizadas estão especificadas em cada seção."
    )
    for run in p_nota.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _calcular_idade(data_nasc_str):
    """Calcula idade em anos a partir de string YYYY-MM-DD."""
    if not data_nasc_str:
        return None
    try:
        from datetime import date
        partes = str(data_nasc_str).split("-")
        nasc = date(int(partes[0]), int(partes[1]), int(partes[2]))
        hoje = date.today()
        return hoje.year - nasc.year - ((hoje.month, hoje.day) < (nasc.month, nasc.day))
    except Exception:
        return None


def _extrair_classificacao(codigo, escores):
    """Extrai a classificação principal de cada teste para destaque no laudo."""
    mapa = {
        "RAVLT": "classif_total",
        "BDI2":  "classificacao",
        "BAI":   "classificacao",
        "SRS2":  "classificacao",
        "ETDAH": "apresentacao_sugerida",
    }
    chave = mapa.get(codigo)
    if chave and isinstance(escores, dict):
        return escores.get(chave)
    if codigo in ("FDT",):
        conds = escores.get("condicoes", {})
        if conds:
            return conds.get("alternar", {}).get("classificacao")
    if codigo in ("TMT",):
        return escores.get("classificacao_b")
    return None
